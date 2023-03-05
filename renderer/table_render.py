"""
表格渲染器
设计模式：抽象工厂模式之具体产品、具体工厂
"""
import math
import copy
import re
from queue import Queue
from renderer.abstract_classes import Render
from abc import ABCMeta, abstractmethod
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from base_utils.base_funcs import get_folder_path
import os
from PIL import Image, ImageFont, ImageDraw
from renderer.pic_render import TableOrderPicRender, TableSuccessRatePicRender, TableAnswerSpacePicRender
from renderer.word_utils.word_utils import TableInserter


class DoubleTableRender(Render, metaclass=ABCMeta):
    @abstractmethod
    def get_test_range(self, test_range=None):
        """
        :return: 获取测试范围样式字符串，传入test_range，返回字符串，如果传入的是None且需要显示测试范围，返回True，否则返回False
        """
        pass

    @abstractmethod
    def adjust_sentence_style_order(self):
        """
        检测表格中是否需要渲染以句子形式呈现的item_node
        """
        pass

    @abstractmethod
    def get_opr_time(self):
        """
        :return: 返回操作时间
        """
        pass

    @abstractmethod
    def get_align(self):
        """
        :return: 返回对齐方式，直接返回WD_PARAGRAPH_ALIGNMENT枚举值
        """
        pass

    @abstractmethod
    def get_data(self):
        """
        :return: 返回题目node的队列
        """
        pass

    # 获取答题区样式
    @abstractmethod
    def get_answer_symbol(self):
        """
        :return: 'en1'表示四线三格，'ch1'表示田字格,'()'表示括号,''表示无样式
        """
        pass

    @abstractmethod
    def get_document(self):
        """
        :return: 获取document,表格附着于document
        """
        pass

    def get_rows_num(self):
        # 总题量为len(self.data)每行容纳columns_num个题目，计算行数，计算结果必须为整数
        sentence_row_num = self.adjust_sentence_style_order()
        word_row_num = math.ceil((self.get_data().qsize() - sentence_row_num) / self.get_columns_num()) * 2
        return word_row_num + sentence_row_num

    @abstractmethod
    def get_columns_num(self):
        """
        :return: 获取列数
        """
        pass

    @abstractmethod
    def get_result_pic_loc_info(self):
        """
        获取 success_num/sum_num 图片位置参数
        :return: 返回元组且坐标系为单元格： (图片x坐标,图片y坐标)，如果返回False，表示不需要插入图片
        """
        pass

    @staticmethod
    def get_result_img(success_num, sum_num):  # 生成单个单词的历史结果图片
        font_style = ImageFont.truetype(os.path.join(get_folder_path(), r'renderer\file', 'font', 'arial.ttf'), 60)
        canvas = Image.new('RGBA', (150, 100), (255, 255, 255, 0))
        draw = ImageDraw.Draw(canvas)
        draw.text((0, 0), '%s/%s' % (str(success_num), str(sum_num)), font=font_style, fill=(0, 0, 0, 160))
        canvas = canvas.crop(canvas.getbbox())
        canvas.save(os.path.join(get_folder_path(), r'renderer\file', 'temp', 'test_result.png'))

    def render(self):
        # 获取数据，绘制表格
        data = self.get_data()  # 获取题目node的队列
        document = self.get_document()  # 获取document
        rows_num = self.get_rows_num()  # 获取行数
        columns_num = self.get_columns_num()  # 获取列数
        table = document.add_table(rows=rows_num, cols=columns_num)  # 添加表格
        is_sentence = False  # 当前node是否为句子形式
        shim = 0

        # 设置表格样式
        answer_symbol = self.get_answer_symbol()  # 答题区样式
        test_range = set()  # 记录范围

        # 向表格中添加数据：题目 + 序号
        for i in range(rows_num):
            for j in range(columns_num):
                # 如果是奇数行，添加题目
                if (i + 1) % 2 == 1 or is_sentence == True:
                    # 如果data为空，跳出循环
                    if data.empty():
                        break
                    item_node = data.get()
                    question_form = item_node[1]
                    item_order = item_node[2]
                    item_node = item_node[0]
                    if question_form == 'word':  # 如果是词语形式的题目
                        table.cell(i, j).text = item_node.item_description
                        p = table.cell(i, j).paragraphs[0]
                        # 添加样式
                        p.style = 'basic_arial'
                        p.alignment = self.get_align()[question_form]
                    elif question_form == 'sentence':  # 如果是句子形式的题目
                        if is_sentence == False and j > 0:  # 跳到下下一行
                            shim = 2  # 写句子要跳行
                            # 处理单词的answer_space
                            for k in range(j):
                                if answer_symbol == 'en1':  # 四线三格
                                    p = table.cell(i + 1, k).paragraphs[0]
                                    TableAnswerSpacePicRender(p=p, pic_name='en1.png', height=0.91).render()
                                elif answer_symbol == 'ch1':  # 田字格
                                    ch_nums = len(table.cell(i, k).text.strip().split(' '))  # 获取cells(i-1,j)的文字，并转为字符数量
                                    p = table.cell(i + 1, k).paragraphs[0]
                                    TableAnswerSpacePicRender(p=p, pic_name='ch1_%s.png' % str(ch_nums),
                                                              height=1.2).render()
                                elif answer_symbol == '()':  # 括号
                                    table.cell(i + 1, k).text = '（          ）'
                                    # 添加样式
                                    p = table.cell(i + 1, k).paragraphs[0]
                                    p.style = 'basic'
                                    p.alignment = self.get_align()[question_form]
                                else:
                                    table.cell(i + 1, k).text = ''
                                    p = table.cell(i + 1, k).paragraphs[0]
                                    p.style = 'basic'
                                    p.alignment = self.get_align()[question_form]

                        is_sentence = True
                        description = item_node.item_description
                        answer = item_node.answer
                        # 检查description中有无{}，如果没有，则查找answer子串，忽略大小写，并用{}包裹
                        if not re.search(r'{.*?}', description):
                            # 查找answer子串，忽略大小写
                            answer_index = re.search(answer, description, re.I).span()
                            # 用{}包裹answer子串
                            description = description[:answer_index[0]] + '{' + description[
                                                                                answer_index[0]:answer_index[
                                                                                    1]] + '}' + description[
                                                                                                answer_index[1]:]
                        # 检查description中有无{}，如果有，则循环替换
                        while re.search(r'{.*?}', description):
                            # 提取第一个{}中的内容，不含{}
                            content = re.search(r'{(.*?)}', description).group(1)
                            answer_space = ''
                            for ch in content:
                                if ch == ' ':
                                    answer_space += ' '
                                else:
                                    answer_space += '*'
                            # 替换第一个{}中的内容，去掉{}
                            description = re.sub(r'{.*?}', answer_space, description, count=1)
                        print(description)
                        current_row = table.rows[i + shim]
                        merged_cell = current_row.cells[0]
                        for cell in current_row.cells[1:]:
                            merged_cell._element.merge(cell._element)
                        # 以连续*为分隔符，分割description，保留分隔符
                        pattern = r'(\*+)'
                        description_list = [ch.strip() for ch in re.split(pattern, description)]
                        # 如果第一个元素为空，则删除
                        if description_list[0] == '':
                            description_list.pop(0)
                        p = merged_cell.paragraphs[0]
                        # p.add_run('')
                        for ch in description_list:
                            if ch == '':
                                p.add_run(' ')
                            if ch.find('*') > -1:
                                underline_num = ch.count('*')
                                underline_str = ' ' * underline_num * int(self.get_length())
                                # underline_str = ' ' * underline_num * 10
                                print(underline_num)
                                p.add_run(underline_str).underline = True
                            else:
                                p.add_run(ch)
                        p.add_run('（%s）' % item_node.deputy_item_description)
                        # 添加样式
                        p.style = 'basic_arial_sentence'
                        p.alignment = self.get_align()[question_form]
                    else:
                        raise Exception('item_node[1]值错误')

                    # 添加序号，png格式
                    order_pos_y = {'word': 0.03, 'sentence': 0.36}
                    TableOrderPicRender(p=p, order=item_order, width=0.5, pos_x=0.03,
                                        pos_y=order_pos_y[question_form]).render()

                    # 记录范围
                    test_range.add(item_node.lesson)

                    # 写入数据库
                    item_node.save_to_record(opr_time=self.get_opr_time(), order=item_order)
                    # # 添加样式
                    # p.style = 'basic_arial'
                    # p.alignment = self.get_align()[question_form]

                    # ------------------------------以下均为可选项------------------------------
                    # 添加题目正确次数/总次数图片
                    test_result = self.get_result_pic_loc_info()
                    if test_result[2]:
                        # 生成临时图片
                        self.get_result_img(item_node.num_success, item_node.num_attempts)
                        TableSuccessRatePicRender(p=p, width=0.29, pos_x=test_result[0], pos_y=test_result[1]).render()
                    if is_sentence:
                        break
                # 如果是偶数行，添加答题区样式
                else:
                    # 如果填充的单元格数量超过了题目数量，跳出循环
                    if (math.floor(i / 2 + 1) - 1) * columns_num + j + 1 > self.get_data().qsize():
                        break
                    if answer_symbol == 'en1':  # 四线三格
                        p = table.cell(i, j).paragraphs[0]
                        TableAnswerSpacePicRender(p=p, pic_name='en1.png', height=0.91).render()
                    elif answer_symbol == 'ch1':  # 田字格
                        ch_nums = len(table.cell(i - 1, j).text.strip().split(' '))  # 获取cells(i-1,j)的文字，并转为字符数量
                        p = table.cell(i, j).paragraphs[0]
                        TableAnswerSpacePicRender(p=p, pic_name='ch1_%s.png' % str(ch_nums), height=1.2).render()
                    elif answer_symbol == '()':  # 括号
                        table.cell(i, j).text = '（          ）'
                        # 添加样式
                        p = table.cell(i, j).paragraphs[0]
                        p.style = 'basic'
                        p.alignment = self.get_align()[question_form]
                    else:
                        table.cell(i, j).text = ''
                        p = table.cell(i, j).paragraphs[0]
                        p.style = 'basic'
                        p.alignment = self.get_align()[question_form]
        if self.get_test_range():  # 如果需要添加考察范围
            test_range_list = list(test_range)  # 将test_range转为列表
            test_range_list.sort()  # 排序
            test_range_list = [str(i) for i in test_range_list]  # 将列表中的元素转为字符串
            TableInserter(document, table).insert_paragraph_before(
                self.get_test_range(test_range_list)).style = 'test_range'

    @abstractmethod
    def get_length(self):
        pass


# ZiCi双层表格渲染器
class ZiCiDoubleTableRender(DoubleTableRender, metaclass=ABCMeta):  # 双层表格渲染器，奇数行为题目描述，偶数行为作答区
    def __init__(self, **kwargs):
        """
        :param data:
        :param document:
        :param opr_time:
        """
        self.data = kwargs['data']  # 题目node列表
        self.document = kwargs['document']
        self.opr_time = kwargs['opr_time']

        self.adjust_sentence_style_order()

    def get_length(self):
        return 5

    # 调整node顺序，使以句子形式考察的默写置于列表最后
    def adjust_sentence_style_order(self):
        """
        :return: n，表示最后n个node为句子形式考察的默写
        """
        temp_list = []
        for node in self.data:
            if len(node.answer) > 4:
                temp_list.append(node)
        # 删除self.data中的已经添加到temp_list中的元素
        for node in temp_list:
            self.data.remove(node)
        # 将temp_list中的元素添加到self.data中
        self.data.extend(temp_list)  # 将temp_list中的元素添加到self.data中

        return len(temp_list)

    def get_test_range(self, test_range=None):
        if test_range:  # test_range为列表
            return '考察范围：第%s课' % '、'.join(test_range)
        else:
            return True

    def get_opr_time(self):
        return self.opr_time

    def get_align(self):
        return {'word': WD_PARAGRAPH_ALIGNMENT.CENTER, 'sentence': WD_PARAGRAPH_ALIGNMENT.LEFT}

    def get_data(self):
        # 创建result空队列
        result = Queue()
        # 遍历data列表，添加到队列中，返回队列
        cnt = 0
        for item in self.data:
            cnt += 1
            if len(item.answer) > 4:  # 如果大于4，则为句子
                if item.deputy_item_description is None:
                    item.deputy_item_description = copy.deepcopy(item.item_description)
                    item.item_description = copy.deepcopy(item.answer)

                result.put((item, 'sentence', cnt))
            else:
                result.put((item, 'word', cnt))
        return result

    @abstractmethod
    def get_answer_symbol(self):
        pass

    def get_document(self):
        return self.document

    def get_result_pic_loc_info(self):
        return 0.67, 0.13, self.data[0].user.current_info['ch_zici_test_result']


# Word双层表格渲染器
class WordDoubleTableRender(DoubleTableRender, metaclass=ABCMeta):  # 双层表格渲染器，奇数行为题目描述，偶数行为作答区
    def __init__(self, **kwargs):
        """
        :param data:
        :param document:
        :param opr_time:
        """
        self.data = kwargs['data']  # 题目node列表
        self.document = kwargs['document']
        self.opr_time = kwargs['opr_time']

        self.adjust_sentence_style_order()

    def get_length(self):  # 下划线长度
        return 2

    # 调整node顺序，使以句子形式考察的默写置于列表最后
    def adjust_sentence_style_order(self):
        temp_list = []
        for node in self.data:
            if not re.search('[\u4e00-\u9fa5]', node.item_description):
                temp_list.append(node)
        # 删除self.data中已经添加到temp_list中的node
        for node in temp_list:
            self.data.remove(node)
        # 将temp_list中的元素添加到self.data中
        self.data.extend(temp_list)  # 将temp_list中的元素添加到self.data中
        return len(temp_list)

    def get_test_range(self, test_range=None):
        if test_range:  # test_range为列表
            return '考察范围：第%s单元' % '、'.join(test_range)
        else:
            return True

    def get_opr_time(self):
        return self.opr_time

    def get_align(self):
        return {'word': WD_PARAGRAPH_ALIGNMENT.CENTER, 'sentence': WD_PARAGRAPH_ALIGNMENT.LEFT}

    def get_data(self):
        # 创建result空队列
        result = Queue()
        # 遍历data列表，添加到队列中，返回队列
        cnt = 0
        for item in self.data:
            cnt += 1
            if not re.search('[\u4e00-\u9fa5]', item.item_description):
                result.put((item, 'sentence', cnt))
            else:
                result.put((item, 'word', cnt))
        return result

    @abstractmethod
    def get_answer_symbol(self):
        pass

    def get_document(self):
        return self.document

    def get_result_pic_loc_info(self):
        return 0.67, 0.13, self.data[0].user.current_info['en_word_test_result']


# 具体表格渲染器
# ZiCi表格渲染器，田字格
class ZiCiTianZiCiDoubleTableRender(ZiCiDoubleTableRender):
    def __init__(self, **kwargs):
        """
        :param data:
        :param document:
        :param opr_time:
        """
        super().__init__(data=kwargs['data'], document=kwargs['document'], opr_time=kwargs['opr_time'])

    def get_answer_symbol(self):
        return 'ch1'

    def get_columns_num(self):
        return 3


# ZiCi表格渲染器，括号
class ZiCiParenthesesZiCiDoubleTableRender(ZiCiDoubleTableRender):
    def __init__(self, **kwargs):
        """
        :param data:
        :param document:
        :param opr_time:
        """
        super().__init__(data=kwargs['data'], document=kwargs['document'], opr_time=kwargs['opr_time'])

    def get_answer_symbol(self):
        return '()'

    def get_columns_num(self):
        return 5


# Word表格渲染器，四线三格
class WordFourLineThreeGridDoubleTableRender(WordDoubleTableRender):
    def __init__(self, **kwargs):
        """
        :param data:
        :param document:
        :param opr_time:
        """
        super().__init__(data=kwargs['data'], document=kwargs['document'], opr_time=kwargs['opr_time'])

    def get_answer_symbol(self):
        return 'en1'

    def get_columns_num(self):
        return 3
