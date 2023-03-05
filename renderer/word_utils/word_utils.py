# -*- coding: utf-8 -*-
import docx
import logging
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# 设置日志的输出格式和输出目的地
logging.basicConfig(format='%(levelname)s: %(message)s', level=logging.INFO)


# 删除传入段落
class DeleteParagraph:
    def __init__(self, p):
        """
        :param paragraph: 需要删除的段落
        """
        self.paragraph = p

    def delete(self):
        p = self.paragraph._element
        p.getparent().remove(p)
        self.paragraph._p = self.paragraph._element = None


# 在指定段落前后添加新段落
class ParagraphInserter:
    def __init__(self, doc, paragraph):
        """
        :param doc_path: 需要添加段落的文档，可以使文档地址，也可以是docx.Document对象
        :param paragraph: 原文中已经存在的段落，新段落的位置将以该段落为原点进行定位
        """
        if isinstance(doc, str):
            self.doc_path = doc
            self.document = docx.Document(doc)
        else:
            self.document = doc

        self.paragraph = None
        if isinstance(paragraph, docx.text.paragraph.Paragraph):
            self.paragraph = paragraph
        elif isinstance(paragraph, str):
            self.paragraph = self.find_paragraph_by_prefix(paragraph)
        else:
            logging.error("参数类型错误：ParagraphInserter传入的段落必须是paragraph类型或str类型")

    # 传入前缀字符串，返回该字符串所在的第一个段落
    def find_paragraph_by_prefix(self, prefix):
        """
        在 word 文档中查找以指定字符串开头的 paragraph 对象
        :param prefix:
        :return:
        """
        for p in self.document.paragraphs:
            print(type(p))
            if p.text.startswith(prefix):
                return p
        logging.error("文档中没有以%s开头的段落" % prefix)
        return None

    # 获取paragraph段落后的第n个段落
    def get_next_paragraph(self, n=1):
        next_paragraph = None
        for i, p in enumerate(self.document.paragraphs):
            # 如果找到了第一个段落
            if p.text == self.paragraph.text:
                # 获取后面的第一个段落
                if i + n < len(self.document.paragraphs):
                    next_paragraph = self.document.paragraphs[i + n]
                    break
                else:
                    logging.error("超出文档段落数")
                    break
        return next_paragraph

    # 获取paragraph段落前的第n个段落
    def get_previous_paragraph(self, n=1):
        previous_paragraph = None
        for i, p in enumerate(self.document.paragraphs):
            # 如果找到了第一个段落
            if p.text == self.paragraph.text:
                # 获取前面的第一个段落
                if i - n >= 0:
                    previous_paragraph = self.document.paragraphs[i - n]
                    break
                else:
                    logging.error("超出文档段落数")
                    break
        return previous_paragraph

    # 以paragraph为原点，向前数n个段落，然后向该段落前插入new_paragraph，如果n为0，则在paragraph前插入
    def insert_paragraph_before_n(self, new_paragraph, n=0):
        previous_paragraph = self.get_previous_paragraph(n)
        if previous_paragraph is not None:
            previous_paragraph.insert_paragraph_before(new_paragraph)
        else:
            # 在document第一段前插入新段落
            self.document.paragraphs[0].insert_paragraph_before(new_paragraph)

    # 以paragraph为原点，向后数n个段落，然后向该段落后插入new_paragraph，如果n为0，则在paragraph后插入
    def insert_paragraph_after_n(self, new_paragraph, n=0):
        next_paragraph = self.get_next_paragraph(n + 1)
        if next_paragraph is not None:
            next_paragraph.insert_paragraph_before(new_paragraph)
        else:
            # 在document最后一段后插入新段落
            self.document.add_paragraph(new_paragraph)


# 在指定Table前后添加新段落
class TableInserter:
    def __init__(self, document, table):
        """
            :param document: 需要添加段落的文档，可以使文档地址，也可以是docx.Document对象
            :param table: 原文中已经存在的段落，新段落的位置将以该段落为原点进行定位
        """
        self.document = None
        self.table = table

        if isinstance(document, str):
            self.doc_path = document
            self.document = docx.Document(document)
        else:
            self.document = document

    # 在指定Table前插入新段落，返回新段落对象
    def insert_paragraph_before(self, new_paragraph):
        for item in self.document.element.body:
            # 如果item对应my_table
            if item == self.table._tbl:
                new_p = docx.oxml.shared.OxmlElement('w:p')  # 创建一个新空paragraph对象
                new_p = docx.text.paragraph.Paragraph(new_p, self.document)  # 转换为docx.text.paragraph.Paragraph对象
                new_p.add_run(new_paragraph)
                item.addprevious(new_p._p)  # 将新段落插入到item前
                return new_p


# ————————————测试用例————————————
if __name__ == '__main__':
    # doc_path = r'C:\Users\77828\Desktop\测试.docx'
    # paragraph = '浩浩乎，平沙无垠1，夐不见人2。'
    # pr = ParagraphInserter(doc_path, paragraph)
    # pr.insert_paragraph_after_n('这是新段落3', 3)
    # 测试
    doc = docx.Document(r'C:\Users\77828\Desktop\test.docx')
    my_table = doc.tables[0]
    TableInserter(doc, my_table).insert_paragraph_before_n('这是新段落', 0)
